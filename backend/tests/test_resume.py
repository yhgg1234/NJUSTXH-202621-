"""简历解析模块 单元测试 —— 文本提取、LLM 解析、服务编排"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from app.resume.extractor import ResumeContentExtractor, normalize_text
from app.resume.models import Education, ParsedResume, ProjectExperience, WorkExperience
from app.resume.parser import (
    ResumeLLMParser,
    _build_parsed_resume,
    _sanitize_json_string,
    _strip_json_fence,
)


# =============================================================================
# normalize_text
# =============================================================================

class TestNormalizeText:
    def test_empty_string(self):
        assert normalize_text("") == ""

    def test_whitespace_collapse(self):
        text = "第一行\n\n\n\n第二行"
        result = normalize_text(text)
        assert "\n\n\n\n" not in result

    def test_control_char_removal(self):
        text = "姓名\x00张三\x1f"
        result = normalize_text(text)
        assert "\x00" not in result
        assert "\x1f" not in result

    def test_fullwidth_punctuation(self):
        text = "张三，清华大学：计算机"
        result = normalize_text(text)
        assert "\uff0c" not in result  # ，
        assert "\uff1a" not in result  # ：
        assert result == "张三,清华大学:计算机"


# =============================================================================
# ResumeContentExtractor
# =============================================================================


class TestContentExtractor:
    def test_file_not_found(self):
        extractor = ResumeContentExtractor()
        with pytest.raises(FileNotFoundError):
            extractor.extract("nonexistent.pdf")

    def test_unsupported_extension(self):
        import tempfile

        extractor = ResumeContentExtractor()
        # 创建一个 .txt 文件以通过存在性检查
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"test")
            tmp_path = f.name
        try:
            with pytest.raises(ValueError, match="不支持"):
                extractor.extract(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_extract_docx(self):
        """生成最小 docx 并验证文本提取。"""
        from docx import Document

        extractor = ResumeContentExtractor()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("张三")
            doc.add_paragraph("技能: Python, FastAPI")
            doc.add_paragraph("")
            doc.save(f.name)
            tmp_path = f.name

        try:
            text = extractor.extract(tmp_path)
            assert "张三" in text
            assert "Python" in text
            # 空段落应被过滤
        finally:
            Path(tmp_path).unlink(missing_ok=True)


# =============================================================================
# _strip_json_fence / _sanitize_json_string
# =============================================================================


def test_strip_markdown_fence():
    assert _strip_json_fence("```json\n{}\n```") == "{}"
    assert _strip_json_fence("```\n{}\n```") == "{}"
    assert _strip_json_fence("{}") == "{}"


def test_sanitize_trailing_comma():
    assert _sanitize_json_string('{"a": 1,}') == '{"a": 1}'
    assert _sanitize_json_string('[1, 2,]') == '[1, 2]'


# =============================================================================
# _build_parsed_resume
# =============================================================================


class TestBuildParsedResume:
    def test_minimal_input(self):
        raw = {"name": "李四"}
        result = _build_parsed_resume(raw, "resume.pdf")
        assert result.name == "李四"
        assert result.file_name == "resume.pdf"
        assert result.skills == []
        assert result.education == []
        assert result.confidence == 0.0

    def test_full_input(self):
        raw = {
            "name": "王五",
            "email": "wangwu@example.com",
            "phone": "13900001111",
            "education": [
                {"school": "北京大学", "degree": "本科", "major": "软件工程", "start_date": "2018-09", "end_date": "2022-06"}
            ],
            "work_experience": [
                {"company": "腾讯", "position": "后端开发", "start_date": "2022-07", "end_date": None, "description": "后端服务开发", "achievements": ["优化接口响应时间 30%"]}
            ],
            "projects": [
                {"name": "电商平台", "role": "核心开发", "start_date": "2021-01", "end_date": "2021-12", "description": "微服务架构", "tech_stacks": ["Spring Boot", "Redis"], "achievements": ["GMV 增长 20%"]}
            ],
            "skills": ["Python", "Java", "Kubernetes"],
            "certificates": ["AWS SAA"],
            "languages": ["中文", "英语"],
            "confidence": 0.92,
        }
        result = _build_parsed_resume(raw, "cv.pdf")
        assert result.name == "王五"
        assert result.email == "wangwu@example.com"
        assert result.confidence == 0.92
        assert len(result.skills) == 3
        assert result.skills[0] == "Python"
        assert len(result.education) == 1
        assert result.education[0].school == "北京大学"
        assert len(result.work_experience) == 1
        assert result.work_experience[0].company == "腾讯"
        assert result.work_experience[0].achievements == ["优化接口响应时间 30%"]
        assert len(result.projects) == 1
        assert result.projects[0].name == "电商平台"
        assert result.projects[0].tech_stacks == ["Spring Boot", "Redis"]

    def test_invalid_confidence_clamped(self):
        raw = {"confidence": 2.5}
        result = _build_parsed_resume(raw, "x.pdf")
        assert result.confidence == 1.0

        raw = {"confidence": -0.5}
        result = _build_parsed_resume(raw, "x.pdf")
        assert result.confidence == 0.0

    def test_corrupted_subelement_skipped(self):
        # Education 必填字段缺失时降级为空字符串，不抛异常
        raw = {
            "education": [{"bad": "data"}],
        }
        result = _build_parsed_resume(raw, "x.pdf")
        assert len(result.education) == 1
        assert result.education[0].school == ""  # 降级为空字符串


# =============================================================================
# ResumeLLMParser
# =============================================================================


class TestResumeLLMParser:
    def test_disabled_when_no_key(self):
        parser = ResumeLLMParser(api_url="", api_key="")
        assert not parser.enabled
        result = parser.parse("一些简历文本", "resume.pdf")
        assert result.file_name == "resume.pdf"
        assert result.confidence == 0.0
        assert result.skills == []

    def test_empty_resume_text_skips_llm(self):
        parser = ResumeLLMParser(api_url="http://localhost", api_key="sk-test")
        assert parser.enabled
        result = parser.parse("", "cv.pdf")
        assert result.confidence == 0.0

    def test_valid_json_response(self):
        parser = ResumeLLMParser(api_url="http://localhost", api_key="sk-test")
        assert parser.enabled
        result = parser._parse_response('{"name":"赵六","skills":["Go"],"confidence":0.88}', "cv.pdf")
        assert result.name == "赵六"
        assert result.skills == ["Go"]
        assert result.confidence == 0.88

    def test_json_with_markdown_fence(self):
        parser = ResumeLLMParser(api_url="http://localhost", api_key="sk-test")
        result = parser._parse_response('```json\n{"name":"钱七"}\n```', "cv.pdf")
        assert result.name == "钱七"

    def test_invalid_json_fallback(self):
        parser = ResumeLLMParser(api_url="http://localhost", api_key="sk-test")
        result = parser._parse_response("这不是 JSON", "cv.pdf")
        assert result.confidence == 0.0

    def test_non_dict_json_fallback(self):
        parser = ResumeLLMParser(api_url="http://localhost", api_key="sk-test")
        result = parser._parse_response("[1, 2, 3]", "cv.pdf")
        assert result.confidence == 0.0


# =============================================================================
# HTTP endpoint integration (light smoke)
# =============================================================================


class TestResumeEndpoints:
    @pytest.fixture(autouse=True)
    def mock_service(self):
        """所有路由测试自动 patch 服务，避免依赖 MongoDB。"""
        with patch("app.routers.resume.ResumeParsingService") as mock_cls:
            mock_inst = mock_cls.return_value
            mock_inst.parse_single.return_value = ParsedResume(
                id="test-001",
                file_name="test.pdf",
                name="测试",
                skills=["Python"],
                confidence=0.9,
            )
            mock_inst.list_resumes.return_value = (
                [ParsedResume(id="test-001", file_name="test.pdf", name="测试")],
                1,
            )
            mock_inst.get_resume.return_value = ParsedResume(
                id="test-001",
                file_name="test.pdf",
                name="测试",
            )
            mock_inst.delete_resume.return_value = True
            mock_inst.search_resumes.return_value = (
                [ParsedResume(id="test-001", file_name="test.pdf", name="测试")],
                1,
            )
            yield mock_inst

    def test_list_resumes(self):
        from fastapi.testclient import TestClient

        from app.main import app

        client = TestClient(app)
        response = client.get("/api/resume/")
        assert response.status_code == 200
        data = response.json()
        assert "items" in data
        assert data["total"] == 1

    def test_get_resume_not_found(self):
        from fastapi.testclient import TestClient

        from app.main import app

        # 覆盖 mock：get_resume 返回 None
        with patch("app.routers.resume._get_service") as svc_mock:
            svc_mock.return_value.get_resume.return_value = None
            client = TestClient(app)
            response = client.get("/api/resume/nonexistent")
            assert response.status_code == 404

    def test_delete_resume(self):
        from fastapi.testclient import TestClient

        from app.main import app

        client = TestClient(app)
        response = client.delete("/api/resume/test-001")
        assert response.status_code == 200
        data = response.json()
        assert data["message"] == "删除成功"

    def test_delete_resume_not_found(self):
        from fastapi.testclient import TestClient

        from app.main import app

        with patch("app.routers.resume._get_service") as svc_mock:
            svc_mock.return_value.delete_resume.return_value = False
            client = TestClient(app)
            response = client.delete("/api/resume/nonexistent")
            assert response.status_code == 404

    def test_search_resumes(self):
        from fastapi.testclient import TestClient

        from app.main import app

        client = TestClient(app)
        response = client.get("/api/resume/search?keyword=Python")
        assert response.status_code == 200
        data = response.json()
        assert data["total"] == 1